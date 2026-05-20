"""
文档解析器 — 支持 PDF（含扫描版OCR via Tesseract）、DOCX、TXT
"""

import os
from config.settings import SUPPORTED_FORMATS, UPLOAD_DIR
from utils.logger import get_logger

logger = get_logger(__name__)

_tesseract_ok = None  # None=未检测, True=可用, False=不可用


def _check_tesseract():
    """检测 Tesseract 是否可用（只检测一次）"""
    global _tesseract_ok
    if _tesseract_ok is not None:
        return _tesseract_ok

    tesseract_path = r"C:\Program Files\Tesseract-OCR\tesseract.exe"
    if not os.path.exists(tesseract_path):
        logger.warning("Tesseract 未安装，扫描版PDF将无法OCR识别")
        _tesseract_ok = False
        return False

    try:
        import pytesseract
        pytesseract.pytesseract.tesseract_cmd = tesseract_path
        # 项目本地中文语言包
        local_tessdata = os.path.join(os.path.dirname(os.path.dirname(__file__)), "data", "tessdata")
        if os.path.exists(os.path.join(local_tessdata, "chi_sim.traineddata")):
            os.environ["TESSDATA_PREFIX"] = local_tessdata
        _tesseract_ok = True
        logger.info("Tesseract OCR 就绪")
        return True
    except ImportError:
        logger.warning("pytesseract 未安装: pip install pytesseract")
        _tesseract_ok = False
        return False
    except Exception as e:
        logger.warning(f"Tesseract 配置失败: {e}")
        _tesseract_ok = False
        return False


def parse_file(file_path: str, force_ocr: bool = False) -> str:
    """根据文件类型自动选择解析器"""
    if not os.path.exists(file_path):
        raise FileNotFoundError(f"文件不存在: {file_path}")
    ext = os.path.splitext(file_path)[1].lower().lstrip(".")
    if ext not in SUPPORTED_FORMATS:
        raise ValueError(f"不支持的文件格式: .{ext}")
    if ext == "pdf":
        return parse_pdf(file_path, force_ocr=force_ocr)
    elif ext == "docx":
        return parse_docx(file_path)
    elif ext == "txt":
        return parse_txt(file_path)
    else:
        raise ValueError(f"未实现解析器: {ext}")


def parse_pdf(file_path: str, force_ocr: bool = False) -> str:
    """PDF解析：文字型用 pdfplumber，扫描版自动 OCR"""
    import pdfplumber

    if force_ocr:
        return _ocr_pdf(file_path)

    all_text = []
    total_chars = 0
    with pdfplumber.open(file_path) as pdf:
        page_count = len(pdf.pages)
        for page in pdf.pages:
            text = page.extract_text()
            if text:
                all_text.append(text)
                total_chars += len(text)
            tables = page.extract_tables()
            for table in tables:
                for row in table:
                    if row:
                        row_text = " | ".join(str(cell) for cell in row if cell)
                        all_text.append(row_text)
                        total_chars += len(row_text)

    avg_chars = total_chars / max(page_count, 1)
    if avg_chars < 30 and page_count > 1:
        logger.info(f"检测到扫描版PDF (平均{avg_chars:.0f}字/页)，切换OCR")
        return _ocr_pdf(file_path)

    logger.info(f"PDF文字提取: {page_count}页, {total_chars}字")
    return "\n".join(all_text)


def _ocr_pdf(file_path: str) -> str:
    """Tesseract OCR 扫描版 PDF：200DPI + 并行 + 加速参数"""
    if not _check_tesseract():
        return (
            "此PDF为扫描版（图片型），需要安装Tesseract OCR才能识别。\n"
            "1. 运行: winget install UB-Mannheim.TesseractOCR\n"
            "2. 下载中文语言包到 data/tessdata/chi_sim.traineddata\n"
            "3. pip install pytesseract"
        )

    import fitz
    import pytesseract
    from PIL import Image
    from concurrent.futures import ThreadPoolExecutor, as_completed

    doc = fitz.open(file_path)
    page_count = doc.page_count

    # Tesseract 加速参数: OEM 1 = LSTM only, PSM 6 = uniform text block
    tesseract_config = '--oem 1 --psm 6'

    def ocr_page(page_num):
        try:
            page = doc.load_page(page_num)
            # 200 DPI 平衡速度与精度
            pix = page.get_pixmap(dpi=200)
            # 转灰度减少处理量
            img = Image.frombytes('RGB', [pix.width, pix.height], pix.samples)
            img = img.convert('L')  # 灰度
            text = pytesseract.image_to_string(img, lang='chi_sim+eng', config=tesseract_config)
            return (page_num, text.strip() if text else '')
        except Exception as e:
            logger.warning(f"OCR第{page_num+1}页失败: {e}")
            return (page_num, '')

    # 并行处理（最多4线程）
    results = {}
    with ThreadPoolExecutor(max_workers=4) as pool:
        futures = {pool.submit(ocr_page, p): p for p in range(page_count)}
        for future in as_completed(futures):
            page_num, text = future.result()
            results[page_num] = text

    doc.close()

    all_text = [results[p] for p in sorted(results.keys()) if results[p]]
    full_text = "\n".join(all_text)
    logger.info(f"OCR完成: {page_count}页, {len(full_text)}字")
    return full_text if full_text.strip() else "OCR未能识别到文字内容"


def parse_docx(file_path: str) -> str:
    """DOCX解析"""
    from docx import Document
    doc = Document(file_path)
    all_text = []
    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            if para.style.name.startswith("Heading"):
                level = para.style.name.replace("Heading ", "")
                prefix = "#" * int(level) if level.isdigit() else "#"
                all_text.append(f"\n{prefix} {text}\n")
            else:
                all_text.append(text)
    for table in doc.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            all_text.append(" | ".join(cells))
    logger.info(f"DOCX解析完成: {file_path}")
    return "\n".join(all_text)


def parse_txt(file_path: str) -> str:
    """TXT解析，自动检测编码"""
    for encoding in ["utf-8", "gbk", "gb2312", "gb18030", "latin-1"]:
        try:
            with open(file_path, "r", encoding=encoding) as f:
                content = f.read()
            logger.info(f"TXT解析完成: {file_path}, 编码={encoding}")
            return content
        except (UnicodeDecodeError, UnicodeError):
            continue
    raise ValueError(f"无法识别文件编码: {file_path}")


def save_uploaded_file(uploaded_file) -> str:
    """保存上传文件"""
    os.makedirs(UPLOAD_DIR, exist_ok=True)
    import time
    original_name = uploaded_file.name
    safe_name = f"{int(time.time())}_{original_name}"
    save_path = os.path.join(UPLOAD_DIR, safe_name)
    with open(save_path, "wb") as f:
        f.write(uploaded_file.getbuffer())
    logger.info(f"文件已保存: {save_path}")
    return save_path


def get_file_info(file_path: str) -> dict:
    """获取文件基本信息"""
    ext = os.path.splitext(file_path)[1].lower().lstrip(".")
    size_mb = os.path.getsize(file_path) / (1024 * 1024)
    return {
        "file_path": file_path,
        "file_type": ext,
        "file_size_mb": round(size_mb, 2),
        "file_name": os.path.basename(file_path),
    }
